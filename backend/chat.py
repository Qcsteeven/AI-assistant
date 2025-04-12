# backend/app.py
from typing import List, Optional, Tuple
import faiss
import numpy as np
from docx import Document
from openai import OpenAI
import os

class VectorBackend:
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance.__init__()
        return cls._instance
    
    def __init__(self):
        if not hasattr(self, 'is_initialized'):  # Защита от повторной инициализации
            self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            self.index: Optional[faiss.Index] = None
            self.embeddings: List[List[float]] = []
            self.chunks: List[str] = []
            self.is_initialized = True

    def load_first_docx(self, chunk_size: int = 500) -> List[str]:
        """Загружает первый .docx файл из папки data"""
        data_dir = "/app/data"  # Путь в контейнере
        if not os.path.exists(data_dir):
            raise FileNotFoundError(f"Папка {data_dir} не существует")

        # Ищем первый .docx файл
        docx_files = [f for f in os.listdir(data_dir) if f.lower().endswith('.docx')]
        if not docx_files:
            raise FileNotFoundError(f"В папке {data_dir} нет .docx файлов")

        first_doc = os.path.join(data_dir, docx_files[0])
        return self.load_docx_chunks(first_doc, chunk_size)

    def load_docx_chunks(self, path: str, chunk_size: int = 500) -> List[str]:
        """Загрузка и разделение документа на чанки"""
        doc = Document(path)
        chunks = []
        
        # Обработка абзацев
        for para in doc.paragraphs:
            if text := para.text.strip():
                chunks.append(text)

        # Обработка таблиц
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows, 1):
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(f"Таблица {table_idx}, строка {row_idx}: {row_text}")

        # Разделение длинных чанков
        final_chunks = []
        for chunk in chunks:
            words = chunk.split()
            if len(words) <= chunk_size:
                final_chunks.append(chunk)
            else:
                final_chunks.extend([
                    " ".join(words[i:i + chunk_size]) 
                    for i in range(0, len(words), chunk_size)
                ])
        
        self.chunks = final_chunks
        return final_chunks

    def get_embedding(self, text: str) -> List[float]:
        """Получение эмбеддинга текста через OpenAI"""
        response = self.client.embeddings.create(
            input=[text],
            model="text-embedding-ada-002"
        )
        return response.data[0].embedding

    def build_faiss_index(self, chunks: List[str]) -> Tuple[faiss.Index, List[List[float]]]:
        """Построение FAISS индекса"""
        self.embeddings = [self.get_embedding(chunk) for chunk in chunks]
        dim = len(self.embeddings[0])
        self.index = faiss.IndexFlatL2(dim)
        self.index.add(np.array(self.embeddings).astype("float32"))
        return self.index, self.embeddings

    def search_similar_chunks(self, query: str, k: int = 3) -> List[str]:
        """Поиск похожих чанков"""
        if not self.index:
            raise ValueError("Индекс не инициализирован. Сначала вызовите build_faiss_index()")
            
        query_vec = np.array([self.get_embedding(query)]).astype("float32")
        _, indices = self.index.search(query_vec, k)
        return [self.chunks[i] for i in indices[0]]

    def chat(self, memory_limit: int = 5):
        """Интерактивный чат с документом"""
        if not self.index:
            raise ValueError("Индекс не инициализирован")
            
        print("💬 Чат с ИИ. Напиши 'exit' для выхода.")
        history = []
        
        while True:
            query = input("\nТы: ")
            if query.lower() == "exit":
                break
                
            top_chunks = self.search_similar_chunks(query)
            context = "\n---\n".join(top_chunks)
            
            history_context = "\n".join(
                f"Вопрос: {q}\nОтвет: {a}" 
                for q, a in history[-memory_limit:]
            )
            
            prompt = (
                "Ты помощник, который отвечает по документу.\n"
                f"История:\n{history_context}\n\n"
                f"Контекст:\n{context}\n\n"
                f"Вопрос: {query}\nОтвет:"
            )
            
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            
            answer = response.choices[0].message.content.strip()
            print("\nGPT:", answer)
            history.append((query, answer))